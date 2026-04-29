"""
Módulo de Lógica Auxiliar — SafetyAI Chat
Responsabilidade: Follow-ups, detecção de intenção de busca e exportação.
"""

import re
import logging
from typing import Optional, List

logger = logging.getLogger(__name__)

# Sugestões rápidas para a tela inicial
SHORTCUT_CHIPS = [
    ("📋", "Como elaborar um PGR?"),
    ("🩺", "Qual a estrutura do PCMSO?"),
    ("⚠️", "O que mudou na NR-1 sobre riscos psicossociais?"),
    ("🔥", "Quais os requisitos da NR-35 para trabalho em altura?"),
    ("🏭", "Como dimensionar a CIPA?"),
    ("📄", "Como fazer uma APR - Análise Preliminar de Risco?"),
    ("🦺", "Quais EPIs obrigatórios para espaço confinado?"),
    ("💰", "Como calcular adicional de insalubridade?"),
]

# Padrões de busca no Google Drive
DRIVE_SEARCH_PATTERNS = re.compile(
    r'\b(?:tem\s+(?:algum|algun[s]?)|(?:me\s+)?(?:mostra|envia|passa|manda)|'
    r'(?:existe[m]?|há|tem)\s+(?:algum|algun[s]?)|encontra|busca|procura|baixar?|'
    r'download\s+d[eo]|arquivo[s]?|documento[s]?|modelo[s]?|planilha|formulário|template)\b',
    re.IGNORECASE,
)

def extract_drive_search_keyword(query: str) -> Optional[str]:
    """Extrai palavra-chave de busca se a query indicar intenção de encontrar documentos."""
    if not DRIVE_SEARCH_PATTERNS.search(query):
        return None
    
    stopwords = {
        'tem', 'algum', 'alguns', 'arquivo', 'arquivos', 'documento', 'documentos',
        'modelo', 'modelos', 'planilha', 'formulário', 'template', 'sobre', 'para',
        'existe', 'existem', 'há', 'mostra', 'envia', 'passa', 'manda', 'encontra',
        'busca', 'procura', 'baixar', 'baixa', 'download', 'de', 'do', 'da', 'um',
        'uma', 'me', 'você', 'voce', 'consegue', 'pode', 'poderia', 'relacionado',
        'relacionados', 'relativo', 'relativos', 'disponível', 'disponivel', 'qualquer',
    }
    tokens = re.findall(r'\b\w[\w\-\.]+\b', query, flags=re.IGNORECASE)
    meaningful = [t for t in tokens if t.lower() not in stopwords and len(t) > 2]
    
    if not meaningful:
        return None
    return " ".join(meaningful[:4])

def generate_follow_ups(query: str, response: str) -> List[str]:
    """Gera 3 sugestões de perguntas baseadas no contexto da conversa."""
    text = (query + " " + response).lower()

    nr_map = {
        "nr-1": ["Quais os prazos para implementação da NR-1?", "Como elaborar o inventário de riscos da NR-1?", "NR-1 exige treinamento de gestão de riscos?"],
        "nr-4": ["Como calcular o quadro do SESMT?", "Quais profissionais compõem o SESMT?", "SESMT é obrigatório para todas as empresas?"],
        "nr-5": ["Como é o processo de eleição da CIPA?", "Quais as atribuições da CIPA?", "Qual a frequência de reuniões da CIPA?"],
        "nr-6": ["Quem é responsável pelo fornecimento de EPI?", "Como documentar o fornecimento de EPIs?", "O que é CA - Certificado de Aprovação de EPI?"],
        "nr-9": ["Qual a diferença entre PGR e PPRA?", "O PPRA foi substituído pelo PGR?", "O que deve conter o inventário de riscos?"],
        "nr-10": ["Quais os requisitos para trabalho em instalações elétricas?", "O que é Prontuário de Instalações Elétricas?", "NR-10 exige treinamento de quantas horas?"],
        "nr-12": ["Quais as proteções obrigatórias em máquinas?", "O que é distância de segurança em máquinas?", "NR-12 exige laudo de conformidade?"],
        "nr-15": ["Como calcular adicional de periculosidade?", "Quais atividades são consideradas insalubres?", "Qual o limite de tolerância para ruído?"],
        "nr-17": ["O que é ergonomia no trabalho?", "NR-17 se aplica a home office?", "Como realizar análise ergonômica do trabalho?"],
        "nr-35": ["O que é Permissão de Trabalho em Altura?", "NR-35 exige qual treinamento?", "Quais EPIs para trabalho em altura?"],
        "pgr": ["Quais os riscos que devem constar no PGR?", "Com que frequência o PGR deve ser revisado?", "O PGR substitui o PPRA e o PCMSO?"],
        "pcmso": ["Com que frequência o PCMSO deve ser revisado?", "Quais exames são obrigatórios no PCMSO?", "O PCMSO precisa de coordenador médico?"],
        "cipa": ["Quais as atribuições do presidente da CIPA?", "Como é o processo de eleição da CIPA?", "O que é SIPAT?"],
        "sesmt": ["Quais os profissionais do SESMT?", "Como dimensionar o SESMT pela NR-4?", "SESMT pode ser terceirizado?"],
        "apr": ["Quais as etapas de uma APR?", "APR é obrigatória por lei?", "Qual a diferença entre APR e PTW?"],
        "epi": ["Como fazer controle de fornecimento de EPI?", "O que é ficha de EPI?", "Qual o prazo de validade do CA?"],
        "insalubridade": ["Quais os graus de insalubridade?", "Como é feita a eliminação de insalubridade?", "Insalubridade e periculosidade podem ser acumulados?"],
        "periculosidade": ["Quais atividades geram adicional de periculosidade?", "Qual o percentual do adicional de periculosidade?", "Como eliminar periculosidade?"],
        "cat": ["Como preencher uma CAT?", "Qual o prazo para emissão da CAT?", "Empresa pode se recusar a emitir CAT?"],
        "espaço confinado": ["O que é a NR-33?", "Quais funções são necessárias em espaço confinado?", "Espaço confinado exige PET?"],
    }

    generic = [
        "Quais as penalidades pelo descumprimento?",
        "Esse tema exige algum documento específico?",
        "Qual NR regulamenta esse assunto?",
        "Quem é responsável por implementar?",
        "Qual o prazo para adequação?",
    ]

    found: List[str] = []
    for kw, questions in nr_map.items():
        if kw in text and len(found) < 3:
            for q in questions:
                if q not in found:
                    found.append(q)
                    if len(found) == 3:
                        break

    while len(found) < 3:
        for q in generic:
            if q not in found:
                found.append(q)
                break

    return found[:3]

def export_chat_docx(messages: list) -> bytes:
    """Gera um arquivo DOCX formatado com o histórico do chat."""
    try:
        from docx import Document
        from docx.shared import Pt, RGBColor
    except ImportError:
        logger.error("Biblioteca 'python-docx' não encontrada.")
        return b""

    doc = Document()
    doc.add_heading('Histórico de Consultoria — SafetyAI', 0)

    for msg in messages:
        role = "Usuário" if msg["role"] == "user" else "SafetyAI"
        p = doc.add_paragraph()
        run = p.add_run(f"{role}: ")
        run.bold = True
        if role == "SafetyAI":
            run.font.color.rgb = RGBColor(34, 197, 94) # Verde neon
        
        p.add_run(str(msg["content"]))
        doc.add_paragraph("-" * 20)

    from io import BytesIO
    target = BytesIO()
    doc.save(target)
    return target.getvalue()
