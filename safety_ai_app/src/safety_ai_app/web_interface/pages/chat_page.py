import streamlit as st
import logging
from typing import Callable, Optional
import markdown
import os
import tempfile
import uuid
import re
import io
from datetime import datetime

from safety_ai_app.theme_config import THEME, _get_material_icon_html
from safety_ai_app.web_interface.shared_styles import inject_glass_styles
from safety_ai_app.input_validators import sanitize_text_input
from safety_ai_app.security.rate_limiter import check_rate_limit, RateLimitExceeded
try:
    from safety_ai_app.feature_access import (
        check_chat_quota,
        increment_chat_messages_today,
        render_chat_quota_exceeded,
        render_upgrade_prompt,
        get_daily_chat_limit,
        get_chat_messages_sent_today,
        is_feature_globally_enabled,
    )
except ImportError:
    logging.getLogger(__name__).error(
        "[chat_page] Módulo feature_access indisponível; aplicando limites conservadores."
    )
    def check_chat_quota(): return False  # noqa: E301
    def increment_chat_messages_today(): pass  # noqa: E301
    def render_chat_quota_exceeded():  # noqa: E301
        st.error("⚠️ Módulo de controle de acesso indisponível. Envio de mensagens bloqueado.")
    def render_upgrade_prompt(feature_label="este recurso"):  # noqa: E301
        st.error("⚠️ Módulo de controle de acesso indisponível. Acesso bloqueado.")
    def get_daily_chat_limit(): return 0  # noqa: E301
    def get_chat_messages_sent_today(): return 0  # noqa: E301
    def is_feature_globally_enabled(feature): return False  # noqa: E301
from safety_ai_app.security.security_logger import log_security_event, SecurityEvent
from safety_ai_app.google_drive_integrator import (
    list_drive_folders,
    get_file_bytes_for_download,
    get_processable_drive_files_in_folder,
    get_file_bytes_by_id,
    list_drive_files_by_keyword,
)
from safety_ai_app.nr_rag_qa import NRQuestionAnswering, make_streamlit_status_callback, is_warmup_complete
from safety_ai_app.text_extractors import get_text_from_file_path

logger = logging.getLogger(__name__)

PROCESSABLE_MIME_TYPES_FOR_RAG = [
    'application/pdf',
    'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
    'text/plain',
    'application/vnd.google-apps.document',
]

MIME_TYPE_DISPLAY = {
    'application/pdf': 'PDF',
    'application/vnd.openxmlformats-officedocument.wordprocessingml.document': 'DOCX',
    'text/plain': 'TXT',
    'application/vnd.google-apps.document': 'Docs',
    'default': 'Arquivo'
}


ALLOWED_TAGS = frozenset({'p', 'br', 'strong', 'em', 'b', 'i', 'u', 'code', 'pre', 'blockquote', 
                          'ul', 'ol', 'li', 'h1', 'h2', 'h3', 'h4', 'h5', 'h6', 'a', 'table', 
                          'thead', 'tbody', 'tr', 'th', 'td', 'hr', 'span', 'div', 'sup', 'sub', 'del', 's'})
ALLOWED_ATTRS = frozenset({'href', 'target', 'rel', 'class', 'id', 'colspan', 'rowspan'})
DANGEROUS_URL_SCHEMES = ('javascript', 'vbscript', 'data', 'file', 'blob')

def _decode_html_entities(text: str) -> str:
    import html
    decoded = html.unescape(text)
    decoded = re.sub(r'%([0-9a-fA-F]{2})', lambda m: chr(int(m.group(1), 16)), decoded)
    decoded = re.sub(r'\\x([0-9a-fA-F]{2})', lambda m: chr(int(m.group(1), 16)), decoded)
    decoded = re.sub(r'\\u([0-9a-fA-F]{4})', lambda m: chr(int(m.group(1), 16)), decoded)
    return decoded

def _is_dangerous_url(url: str) -> bool:
    decoded = _decode_html_entities(url)
    normalized = re.sub(r'\s+', '', decoded.lower())
    for scheme in DANGEROUS_URL_SCHEMES:
        if normalized.startswith(f'{scheme}:'):
            return True
        if re.match(rf'^{re.escape(scheme)}\s*:', normalized):
            return True
    return False

def _sanitize_tag(match: re.Match) -> str:
    full_tag = match.group(0)
    tag_name = match.group(1).lower() if match.group(1) else match.group(2).lower()
    
    if tag_name not in ALLOWED_TAGS:
        return ''
    
    is_closing = full_tag.startswith('</')
    is_self_closing = full_tag.rstrip().endswith('/>')
    
    if is_closing:
        return f'</{tag_name}>'
    
    safe_attrs = []
    attr_pattern = re.compile(r'(\w+)\s*=\s*(?:"([^"]*)"|\'([^\']*)\'|(\S+))', re.IGNORECASE)
    for attr_match in attr_pattern.finditer(full_tag):
        attr_name = attr_match.group(1).lower()
        attr_value = attr_match.group(2) or attr_match.group(3) or attr_match.group(4) or ''
        
        if attr_name.startswith('on'):
            continue
        
        if attr_name in ('href', 'src', 'action', 'formaction'):
            if _is_dangerous_url(attr_value):
                continue
        
        if attr_name in ALLOWED_ATTRS:
            escaped_value = attr_value.replace('"', '&quot;').replace('<', '&lt;').replace('>', '&gt;')
            safe_attrs.append(f'{attr_name}="{escaped_value}"')
    
    attrs_str = ' '.join(safe_attrs)
    if attrs_str:
        return f'<{tag_name} {attrs_str}{"/" if is_self_closing else ""}>'
    return f'<{tag_name}{"/" if is_self_closing else ""}>'

def _sanitize_html(html: str) -> str:
    html = re.sub(r'<!--.*?-->', '', html, flags=re.DOTALL)
    html = re.sub(r'<\s*/?\s*(\w+)[^>]*/?>', _sanitize_tag, html)
    html = re.sub(r'\{\{[^}]*\}\}|\[\[[^\]]*\]\]', '', html)
    return html

def _safe_markdown(text) -> str:
    if not text:
        return ""
    if isinstance(text, dict):
        text = text.get("content", text.get("answer", str(text)))
    if not isinstance(text, str):
        text = str(text)
    
    safe_text = text.replace('<', '&lt;').replace('>', '&gt;').replace('"', '&quot;').replace("'", '&#39;')
    
    try:
        html = markdown.markdown(text, extensions=['extra', 'nl2br'], output_format='html')
        html = _sanitize_html(html)
        html = re.sub(r'<a href="([^"]*)"([^>]*)>', r'<a href="\1" target="_blank" rel="noopener noreferrer"\2>', html)
        return html
    except Exception as e:
        logger.error(f"Erro markdown: {e}")
        return safe_text


def _inject_chat_styles():
    st.markdown("""
    <style>
        /* === CYBER-NEON CHAT - REDESIGN INOVADOR === */
        :root {
            --neon-green: #4ADE80;
            --neon-green-dark: #22C55E;
            --neon-glow: rgba(74, 222, 128, 0.4);
            --bg-void: #020617;
            --bg-space: #0B1220;
            --bg-nebula: #0F172A;
            --glass-bg: rgba(15, 23, 42, 0.85);
            --glass-border: rgba(74, 222, 128, 0.15);
            --text-primary: #F8FAFC;
            --text-secondary: #94A3B8;
            --text-muted: #64748B;
        }

        /* === ANIMAÇÕES === */
        @keyframes fadeInUp {
            from { opacity: 0; transform: translateY(15px); }
            to { opacity: 1; transform: translateY(0); }
        }
        @keyframes pulse {
            0%, 100% { opacity: 0.4; }
            50% { opacity: 1; }
        }
        @keyframes glow {
            0%, 100% { box-shadow: 0 0 5px var(--neon-glow); }
            50% { box-shadow: 0 0 20px var(--neon-glow), 0 0 30px var(--neon-glow); }
        }

        /* === CONTAINER PRINCIPAL === */
        .chat-main-container {
            display: flex;
            flex-direction: column;
            height: calc(100vh - 200px);
            min-height: 500px;
        }

        /* === HEADER DO CHAT === */
        .chat-header {
            display: flex;
            align-items: center;
            gap: 12px;
            padding: 16px 20px;
            background: linear-gradient(135deg, var(--bg-nebula) 0%, var(--bg-space) 100%);
            border: 1px solid var(--glass-border);
            border-radius: 16px;
            margin-bottom: 16px;
        }
        .chat-header-icon {
            width: 48px;
            height: 48px;
            border-radius: 12px;
            background: linear-gradient(135deg, var(--neon-green) 0%, var(--neon-green-dark) 100%);
            display: flex;
            align-items: center;
            justify-content: center;
            animation: glow 3s ease-in-out infinite;
        }
        .chat-header-icon svg { width: 28px; height: 28px; color: #000; }
        .chat-header-text h1 {
            margin: 0;
            font-size: 1.5rem;
            font-weight: 700;
            color: var(--text-primary);
            font-family: 'Orbitron', sans-serif;
        }
        .chat-header-text p {
            margin: 4px 0 0 0;
            font-size: 0.85rem;
            color: var(--text-muted);
        }

        /* === ÁREA DE MENSAGENS === */
        .messages-container {
            flex: 1;
            overflow-y: auto;
            padding: 16px;
            background: var(--bg-void);
            border: 1px solid var(--glass-border);
            border-radius: 16px;
            margin-bottom: 16px;
        }

        /* === MENSAGEM DO USUÁRIO === */
        .msg-user {
            display: flex;
            justify-content: flex-end;
            margin-bottom: 16px;
            animation: fadeInUp 0.3s ease-out;
        }
        .msg-user-bubble {
            max-width: 75%;
            padding: 14px 18px;
            background: linear-gradient(135deg, #166534 0%, #15803D 100%);
            color: #fff;
            border-radius: 18px 18px 4px 18px;
            font-size: 0.95rem;
            line-height: 1.5;
            box-shadow: 0 4px 12px rgba(22, 101, 52, 0.3);
        }

        /* === MENSAGEM DA IA === */
        .msg-ai {
            display: flex;
            gap: 12px;
            margin-bottom: 20px;
            animation: fadeInUp 0.4s ease-out;
        }
        .msg-ai-avatar {
            flex-shrink: 0;
            width: 40px;
            height: 40px;
            border-radius: 12px;
            background: linear-gradient(135deg, var(--neon-green) 0%, var(--neon-green-dark) 100%);
            display: flex;
            align-items: center;
            justify-content: center;
            box-shadow: 0 0 15px var(--neon-glow);
        }
        .msg-ai-avatar svg { width: 22px; height: 22px; color: #000; }
        .msg-ai-content {
            flex: 1;
            padding: 16px 20px;
            background: var(--glass-bg);
            border: 1px solid var(--glass-border);
            border-radius: 4px 18px 18px 18px;
            color: var(--text-primary);
            font-size: 0.95rem;
            line-height: 1.6;
            backdrop-filter: blur(10px);
        }
        .msg-ai-content p { margin: 0 0 12px 0; }
        .msg-ai-content p:last-child { margin-bottom: 0; }
        .msg-ai-content strong { color: var(--neon-green); }
        .msg-ai-content em { color: var(--text-secondary); font-style: italic; }
        .msg-ai-content a { color: var(--neon-green); text-decoration: underline; }
        .msg-ai-content ul, .msg-ai-content ol { margin: 8px 0; padding-left: 20px; }
        .msg-ai-content li { margin-bottom: 4px; }

        /* === INDICADOR DE DIGITAÇÃO === */
        .typing-indicator {
            display: flex;
            gap: 12px;
            align-items: center;
            margin-bottom: 16px;
            animation: fadeInUp 0.3s ease-out;
        }
        .typing-dots {
            display: flex;
            gap: 6px;
            padding: 12px 16px;
            background: var(--glass-bg);
            border: 1px solid var(--glass-border);
            border-radius: 18px;
        }
        .typing-dots span {
            width: 8px;
            height: 8px;
            background: var(--neon-green);
            border-radius: 50%;
            animation: pulse 1.4s infinite;
        }
        .typing-dots span:nth-child(2) { animation-delay: 0.2s; }
        .typing-dots span:nth-child(3) { animation-delay: 0.4s; }

        /* === TELA DE BOAS-VINDAS === */
        .welcome-screen {
            display: flex;
            flex-direction: column;
            align-items: center;
            justify-content: center;
            padding: 40px 20px;
            text-align: center;
        }
        .welcome-icon {
            width: 80px;
            height: 80px;
            border-radius: 20px;
            background: linear-gradient(135deg, var(--neon-green) 0%, var(--neon-green-dark) 100%);
            display: flex;
            align-items: center;
            justify-content: center;
            margin-bottom: 24px;
            animation: glow 3s ease-in-out infinite;
        }
        .welcome-icon svg { width: 44px; height: 44px; color: #000; }
        .welcome-title {
            font-size: 1.4rem;
            font-weight: 600;
            color: var(--text-primary);
            margin-bottom: 8px;
        }
        .welcome-subtitle {
            font-size: 0.95rem;
            color: var(--text-secondary);
            max-width: 400px;
            line-height: 1.5;
            margin-bottom: 24px;
        }

        /* === SUGESTÕES DE PERGUNTAS === */
        .suggestions-container {
            display: flex;
            flex-wrap: wrap;
            gap: 10px;
            justify-content: center;
            max-width: 600px;
        }
        .suggestion-chip {
            padding: 10px 16px;
            background: var(--glass-bg);
            border: 1px solid var(--glass-border);
            border-radius: 20px;
            color: var(--text-primary);
            font-size: 0.85rem;
            cursor: pointer;
            transition: all 0.2s ease;
        }
        .suggestion-chip:hover {
            background: rgba(74, 222, 128, 0.15);
            border-color: var(--neon-green);
            transform: translateY(-2px);
        }

        /* === ÁREA DE INPUT === */
        .input-container {
            background: var(--glass-bg);
            border: 1px solid var(--glass-border);
            border-radius: 16px;
            padding: 12px 16px;
        }
        .input-row {
            display: flex;
            gap: 12px;
            align-items: center;
        }

        /* === DRAWER DE FERRAMENTAS === */
        .tools-drawer {
            background: var(--glass-bg);
            border: 1px solid var(--glass-border);
            border-radius: 12px;
            padding: 16px;
            margin-bottom: 16px;
        }
        .tools-drawer-header {
            display: flex;
            align-items: center;
            gap: 8px;
            color: var(--neon-green);
            font-weight: 600;
            margin-bottom: 12px;
        }
        .tools-drawer-header svg { width: 18px; height: 18px; }

        /* === ARQUIVOS ATIVOS === */
        .active-files {
            display: flex;
            flex-wrap: wrap;
            gap: 8px;
            padding: 12px 0;
        }
        .file-chip {
            display: inline-flex;
            align-items: center;
            gap: 6px;
            padding: 6px 12px;
            background: rgba(74, 222, 128, 0.1);
            border: 1px solid var(--glass-border);
            border-radius: 16px;
            font-size: 0.8rem;
            color: var(--text-primary);
        }
        .file-chip svg { width: 14px; height: 14px; color: var(--neon-green); }

        /* === RESPONSIVIDADE === */
        @media (max-width: 768px) {
            .chat-header { padding: 12px 16px; }
            .chat-header-icon { width: 40px; height: 40px; }
            .chat-header-text h1 { font-size: 1.2rem; }
            .msg-user-bubble, .msg-ai-content { max-width: 90%; }
            .welcome-icon { width: 60px; height: 60px; }
            .welcome-title { font-size: 1.2rem; }
        }
    </style>
    """, unsafe_allow_html=True)


_SHORTCUT_CHIPS = [
    ("📋", "Como elaborar um PGR?"),
    ("🩺", "Qual a estrutura do PCMSO?"),
    ("⚠️", "O que mudou na NR-1 sobre riscos psicossociais?"),
    ("🔥", "Quais os requisitos da NR-35 para trabalho em altura?"),
    ("🏭", "Como dimensionar a CIPA?"),
    ("📄", "Como fazer uma APR - Análise Preliminar de Risco?"),
    ("🦺", "Quais EPIs obrigatórios para espaço confinado?"),
    ("💰", "Como calcular adicional de insalubridade?"),
]


def _render_welcome_screen():
    st.markdown(f"""
    <div class="welcome-screen">
        <div class="welcome-icon">
            {_get_material_icon_html("smart_toy")}
        </div>
        <div class="welcome-title">Olá! Como posso ajudar?</div>
        <div class="welcome-subtitle">
            Sou seu assistente especializado em Saúde e Segurança do Trabalho.<br>
            Pergunte sobre NRs, dimensionamento, EPIs, documentos e muito mais.
        </div>
        <div class="suggestions-container">
    """, unsafe_allow_html=True)

    row1 = _SHORTCUT_CHIPS[:4]
    row2 = _SHORTCUT_CHIPS[4:]
    cols1 = st.columns(4)
    for i, (icon, label) in enumerate(row1):
        with cols1[i]:
            if st.button(f"{icon} {label}", key=f"shortcut_{i}", use_container_width=True):
                st.session_state.pending_query = label
                st.rerun()
    cols2 = st.columns(4)
    for i, (icon, label) in enumerate(row2):
        with cols2[i]:
            if st.button(f"{icon} {label}", key=f"shortcut_{i+4}", use_container_width=True):
                st.session_state.pending_query = label
                st.rerun()

    st.markdown("</div></div>", unsafe_allow_html=True)


def _generate_follow_ups(query: str, response: str) -> list[str]:
    """Generate 3 follow-up question suggestions based on the query/response content."""
    text = (query + " " + response).lower()

    nr_map = {
        "nr-1": ["Quais os prazos para implementação da NR-1?", "Como elaborar o inventário de riscos da NR-1?", "NR-1 exige treinamento de gestão de riscos?"],
        "nr-4": ["Como calcular o quadro do SESMT?", "Quais profissionais compõem o SESMT?", "SESMT é obrigatório para todas as empresas?"],
        "nr-5": ["Como é o processo de eleição da CIPA?", "Quais as atribuições da CIPA?", "Qual a frequência de reuniões da CIPA?"],
        "nr-6": ["Quem é responsável pelo fornecimento de EPI?", "Como documentar o fornecimento de EPIs?", "O que é CA - Certificado de Aprovação de EPI?"],
        "nr-9": ["Qual a diferença entre PGR e PPRA?", "O PPRA foi substituído pelo PGR?", "O que deve conter o inventário de riscos?"],
        "nr-10": ["Quais os requisitos para trabalho em instalações elétricas?", "O que é Prontuário de Instalações Elétricas?", "NR-10 exige treinamento de quantas horas?"],
        "nr-12": ["Quais as proteções obrigatórias em máquinas?", "O que é distância de segurança em máquinas?", "NR-12 exige laudo de conformidade?"],
        "nr-15": ["Como calcular adicional de periculosidade?", "Quais atividades são consideradas insalubres?", "Qual o limite de tolerância para ruído?"],
        "nr-17": ["O que é ergonomia no trabalho?", "NR-17 se aplica a home office?", "Como realizar análise ergonômica do trabalho?"],
        "nr-35": ["O que é Permissão de Trabalho em Altura?", "NR-35 exige qual treinamento?", "Quais EPIs para trabalho em altura?"],
        "pgr": ["Quais os riscos que devem constar no PGR?", "Com que frequência o PGR deve ser revisado?", "O PGR substitui o PPRA e o PCMSO?"],
        "pcmso": ["Com que frequência o PCMSO deve ser revisado?", "Quais exames são obrigatórios no PCMSO?", "O PCMSO precisa de coordenador médico?"],
        "cipa": ["Quais as atribuições do presidente da CIPA?", "Como é o processo de eleição da CIPA?", "O que é SIPAT?"],
        "sesmt": ["Quais os profissionais do SESMT?", "Como dimensionar o SESMT pela NR-4?", "SESMT pode ser terceirizado?"],
        "apr": ["Quais as etapas de uma APR?", "APR é obrigatória por lei?", "Qual a diferença entre APR e PTW?"],
        "epi": ["Como fazer controle de fornecimento de EPI?", "O que é ficha de EPI?", "Qual o prazo de validade do CA?"],
        "insalubridade": ["Quais os graus de insalubridade?", "Como é feita a eliminação de insalubridade?", "Insalubridade e periculosidade podem ser acumulados?"],
        "periculosidade": ["Quais atividades geram adicional de periculosidade?", "Qual o percentual do adicional de periculosidade?", "Como eliminar periculosidade?"],
        "cat": ["Como preencher uma CAT?", "Qual o prazo para emissão da CAT?", "Empresa pode se recusar a emitir CAT?"],
        "espaço confinado": ["O que é a NR-33?", "Quais funções são necessárias em espaço confinado?", "Espaço confinado exige PET?"],
        "lgpd": ["LGPD se aplica a dados de saúde ocupacional?", "Como proteger dados do PCMSO?", "Qual o prazo de guarda dos prontuários?"],
    }

    generic = [
        "Quais as penalidades pelo descumprimento?",
        "Esse tema exige algum documento específico?",
        "Qual NR regulamenta esse assunto?",
        "Quem é responsável por implementar?",
        "Qual o prazo para adequação?",
    ]

    found: list[str] = []
    for kw, questions in nr_map.items():
        if kw in text and len(found) < 3:
            for q in questions:
                if q not in found:
                    found.append(q)
                    if len(found) == 3:
                        break

    while len(found) < 3:
        for q in generic:
            if q not in found:
                found.append(q)
                break

    return found[:3]


_DRIVE_SEARCH_PATTERNS = re.compile(
    r'\b(?:tem\s+(?:algum|algun[s]?)|(?:me\s+)?(?:mostra|envia|passa|manda)|'
    r'(?:existe[m]?|há|tem)\s+(?:algum|algun[s]?)|encontra|busca|procura|baixar?|'
    r'download\s+d[eo]|arquivo[s]?|documento[s]?|modelo[s]?|planilha|formulário|template)\b',
    re.IGNORECASE,
)

def _extract_drive_search_keyword(query: str) -> Optional[str]:
    """If the query is asking for a document/file, extract the search keyword. Returns None otherwise."""
    if not _DRIVE_SEARCH_PATTERNS.search(query):
        return None
    stopwords = {
        'tem', 'algum', 'alguns', 'arquivo', 'arquivos', 'documento', 'documentos',
        'modelo', 'modelos', 'planilha', 'formulário', 'template', 'sobre', 'para',
        'existe', 'existem', 'há', 'mostra', 'envia', 'passa', 'manda', 'encontra',
        'busca', 'procura', 'baixar', 'baixa', 'download', 'de', 'do', 'da', 'um',
        'uma', 'me', 'você', 'voce', 'consegue', 'pode', 'poderia', 'relacionado',
        'relacionados', 'relativo', 'relativos', 'disponível', 'disponivel', 'qualquer',
    }
    tokens = re.findall(r'\b\w[\w\-\.]+\b', query, flags=re.IGNORECASE)
    meaningful = [t for t in tokens if t.lower() not in stopwords and len(t) > 2]
    if not meaningful:
        return None
    return " ".join(meaningful[:4])


def _export_chat_docx(messages: list) -> bytes:
    """Generate a formatted DOCX with all chat messages."""
    try:
        from docx import Document
        from docx.shared import Pt, RGBColor, Inches
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    except ImportError:
        return b""

    doc = Document()

    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)

    title = doc.add_heading('SafetyAI — Conversa Exportada', level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.runs[0].font.color.rgb = RGBColor(0x5B, 0x21, 0xB6)

    sub = doc.add_paragraph(f"Exportado em: {datetime.now().strftime('%d/%m/%Y às %H:%M')}")
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.runs[0].font.color.rgb = RGBColor(0x64, 0x74, 0x8B)
    sub.runs[0].font.size = Pt(9)

    doc.add_paragraph()
    doc.add_paragraph("─" * 60)
    doc.add_paragraph()

    for msg in messages:
        role = msg.get("role", "")
        content = msg.get("content", "")
        if isinstance(content, dict):
            content = content.get("answer", content.get("content", str(content)))
        content = str(content) if content else ""

        if role == "user":
            p = doc.add_paragraph()
            p.add_run("Você:  ").bold = True
            r = p.runs[0]
            r.font.color.rgb = RGBColor(0x06, 0xB6, 0xD4)
            r.font.size = Pt(10)
            p.add_run(content)
        else:
            p = doc.add_paragraph()
            p.add_run("SafetyAI:  ").bold = True
            r = p.runs[0]
            r.font.color.rgb = RGBColor(0x4A, 0xDE, 0x80)
            r.font.size = Pt(10)
            clean = re.sub(r'[*_`#>]+', '', content)
            p.add_run(clean)

            downloads = msg.get("suggested_downloads", [])
            if downloads:
                dp = doc.add_paragraph()
                dp.add_run("📎 Documentos relacionados: ").italic = True
                dp.add_run(", ".join(d.get("document_name", "") for d in downloads))

        doc.add_paragraph()

    doc.add_paragraph("─" * 60)
    footer = doc.add_paragraph("Gerado pelo SafetyAI App — Assistente de Saúde e Segurança do Trabalho")
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.runs[0].font.color.rgb = RGBColor(0x64, 0x74, 0x8B)
    footer.runs[0].font.size = Pt(8)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _render_message(msg: dict, idx: int, markdown_func: Callable):
    content = msg.get("content", "")
    if isinstance(content, dict):
        content = content.get("answer", content.get("content", str(content)))
    if not isinstance(content, str):
        content = str(content) if content else ""
    
    if msg.get("role") == "user":
        safe_content = content.replace('<', '&lt;').replace('>', '&gt;')
        st.markdown(f"""
        <div class="msg-user">
            <div class="msg-user-bubble">{safe_content}</div>
        </div>
        """, unsafe_allow_html=True)
    else:
        content_html = markdown_func(content)
        st.markdown(f"""
        <div class="msg-ai">
            <div class="msg-ai-avatar">
                {_get_material_icon_html("smart_toy")}
            </div>
            <div class="msg-ai-content">{content_html}</div>
        </div>
        """, unsafe_allow_html=True)
        
        if "suggested_downloads" in msg and msg["suggested_downloads"]:
            _render_downloads(msg["suggested_downloads"], idx)


def _render_downloads(downloads: list, msg_idx: int):
    mime_icons = {
        'application/pdf': '📕',
        'application/vnd.openxmlformats-officedocument.wordprocessingml.document': '📘',
        'text/plain': '📄',
        'application/vnd.google-apps.document': '📝',
    }
    mime_labels = {
        'application/pdf': 'PDF',
        'application/vnd.openxmlformats-officedocument.wordprocessingml.document': 'DOCX',
        'text/plain': 'TXT',
        'application/vnd.google-apps.document': 'Google Doc',
    }

    st.markdown("""
    <div style="
        background: rgba(139,92,246,0.08);
        border: 1px solid rgba(139,92,246,0.25);
        border-radius: 12px;
        padding: 12px 16px;
        margin: 10px 0 4px 0;
    ">
        <div style="font-size:0.82rem; color:#a78bfa; font-weight:600; margin-bottom:8px; letter-spacing:0.03em;">
            📁 Documentos Relacionados na Biblioteca
        </div>
    </div>
    """, unsafe_allow_html=True)

    cols = st.columns(min(len(downloads), 3))
    for i, doc in enumerate(downloads):
        with cols[i % 3]:
            icon = mime_icons.get(doc.get('file_type', ''), '📄')
            label_type = mime_labels.get(doc.get('file_type', ''), 'Arquivo')
            short_name = doc['document_name'][:22] + "…" if len(doc['document_name']) > 22 else doc['document_name']
            try:
                file_bytes = get_file_bytes_for_download(
                    st.session_state.user_drive_service,
                    doc['drive_file_id'],
                    doc['file_type'],
                    doc['file_type']
                )
                ext = {
                    'application/pdf': '.pdf',
                    'application/vnd.openxmlformats-officedocument.wordprocessingml.document': '.docx',
                    'text/plain': '.txt',
                    'application/vnd.google-apps.document': '.docx'
                }.get(doc['file_type'], '')
                filename = doc['document_name'] + ext if not doc['document_name'].endswith(ext) else doc['document_name']
                st.download_button(
                    label=f"{icon} {short_name} [{label_type}]",
                    data=file_bytes,
                    file_name=filename,
                    mime=doc['file_type'],
                    key=f"dl_{msg_idx}_{i}_{doc['drive_file_id']}",
                    use_container_width=True,
                    help=doc['document_name'],
                )
            except Exception as e:
                logger.error(f"Erro download: {e}")
                st.markdown(f"<span style='font-size:0.78rem;color:#64748b;'>{icon} {short_name}</span>", unsafe_allow_html=True)


def _render_follow_ups(follow_ups: list[str]):
    """Render follow-up question chips below AI response."""
    if not follow_ups:
        return
    st.markdown("""
    <div style="margin: 6px 0 2px 0; font-size:0.78rem; color:#64748b; font-style:italic;">
        Perguntas de acompanhamento:
    </div>
    """, unsafe_allow_html=True)
    cols = st.columns(len(follow_ups))
    for i, q in enumerate(follow_ups):
        with cols[i]:
            if st.button(q, key=f"followup_{id(q)}_{i}", use_container_width=True):
                st.session_state.pending_query = q
                st.session_state.last_follow_ups = []
                st.rerun()


def _render_typing_indicator():
    st.markdown(f"""
    <div class="typing-indicator">
        <div class="msg-ai-avatar">
            {_get_material_icon_html("smart_toy")}
        </div>
        <div class="typing-dots">
            <span></span><span></span><span></span>
        </div>
    </div>
    """, unsafe_allow_html=True)


def render_page(process_markdown_for_external_links_func: Callable[[str], str] | None = None) -> None:
    markdown_func = process_markdown_for_external_links_func if process_markdown_for_external_links_func else _safe_markdown

    if not is_feature_globally_enabled("chat"):
        render_upgrade_prompt("Chat SST")
        return

    inject_glass_styles()
    _inject_chat_styles()
    
    if "messages" not in st.session_state:
        st.session_state.messages = []
    if "active_context_files" not in st.session_state:
        st.session_state.active_context_files = []
    if "show_tools" not in st.session_state:
        st.session_state.show_tools = False
    if "pending_query" not in st.session_state:
        st.session_state.pending_query = None
    if "chat_mode" not in st.session_state:
        st.session_state.chat_mode = "deep"
    if "last_follow_ups" not in st.session_state:
        st.session_state.last_follow_ups = []

    st.markdown(f"""
    <div class="chat-header">
        <div class="chat-header-icon">
            {_get_material_icon_html("smart_toy")}
        </div>
        <div class="chat-header-text">
            <h1>Chat SST</h1>
            <p>Assistente de Segurança do Trabalho</p>
        </div>
    </div>
    """, unsafe_allow_html=True)

    hcols = st.columns([1, 1, 1, 1])
    with hcols[0]:
        quick_active = st.session_state.chat_mode == "quick"
        if st.button(
            "⚡ Consulta Rápida",
            key="mode_quick",
            use_container_width=True,
            type="primary" if quick_active else "secondary",
            help="Respostas objetivas e diretas",
        ):
            st.session_state.chat_mode = "quick"
            st.rerun()
    with hcols[1]:
        deep_active = st.session_state.chat_mode == "deep"
        if st.button(
            "🔬 Análise Técnica",
            key="mode_deep",
            use_container_width=True,
            type="primary" if deep_active else "secondary",
            help="Análise completa com referências normativas",
        ):
            st.session_state.chat_mode = "deep"
            st.rerun()
    with hcols[2]:
        if st.session_state.messages:
            docx_bytes = _export_chat_docx(st.session_state.messages)
            if docx_bytes:
                fname = f"SafetyAI_Chat_{datetime.now().strftime('%Y%m%d_%H%M')}.docx"
                st.download_button(
                    "⬇ Exportar",
                    data=docx_bytes,
                    file_name=fname,
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    use_container_width=True,
                    key="export_chat_docx",
                    help="Baixar conversa como documento Word",
                )
    with hcols[3]:
        if st.session_state.messages:
            if st.button("🗑️ Limpar Chat", key="clear_chat_top", use_container_width=True, help="Apagar todo o histórico"):
                st.session_state.messages = []
                st.session_state.last_follow_ups = []
                st.rerun()

    if not is_warmup_complete():
        st.info(
            "⏳ Os modelos de IA estão sendo carregados em segundo plano. "
            "A primeira resposta pode levar alguns segundos a mais enquanto o carregamento termina.",
            icon=None,
        )

    _daily_limit = get_daily_chat_limit()
    if _daily_limit > 0:
        _sent_today = get_chat_messages_sent_today()
        _remaining = max(0, _daily_limit - _sent_today)
        _pct = int((_sent_today / _daily_limit) * 100)
        _bar_color = "#ef4444" if _remaining == 0 else ("#f59e0b" if _remaining <= 2 else "#4ade80")
        st.markdown(
            f"""
            <div style="
                display:flex; align-items:center; gap:10px;
                background: rgba(15,23,42,0.7);
                border: 1px solid rgba(74,222,128,0.12);
                border-radius:10px; padding:8px 14px; margin-bottom:8px;
                font-size:0.83rem; color:#94a3b8;
            ">
                <span>💬 Mensagens hoje:</span>
                <div style="flex:1; background:#1e293b; border-radius:4px; height:6px; overflow:hidden;">
                    <div style="width:{_pct}%; background:{_bar_color}; height:100%; border-radius:4px; transition:width .3s;"></div>
                </div>
                <span style="color:#f8fafc; font-weight:600;">{_sent_today}/{_daily_limit}</span>
            </div>
            """,
            unsafe_allow_html=True,
        )

    with st.expander("🛠️ Ferramentas Avançadas", expanded=st.session_state.show_tools):
        tab_local, tab_drive = st.tabs(["📁 Upload Local", "☁️ Google Drive"])
        
        with tab_local:
            uploaded_files = st.file_uploader(
                "Envie documentos para contexto",
                type=["pdf", "docx", "txt"],
                accept_multiple_files=True,
                key="local_uploader",
                label_visibility="collapsed"
            )
            if uploaded_files:
                MAX_FILE_SIZE_MB = 15
                new_files = []
                rejected = []
                for f in uploaded_files:
                    size_mb = len(f.getvalue()) / (1024 * 1024)
                    if size_mb > MAX_FILE_SIZE_MB:
                        rejected.append(f.name)
                        log_security_event(
                            SecurityEvent.FILE_REJECTED,
                            file_name=f.name,
                            file_size_mb=size_mb,
                            detail=f"Arquivo excede {MAX_FILE_SIZE_MB} MB",
                            feature="chat_upload",
                        )
                        continue
                    new_files.append({
                        'id': f"local_{uuid.uuid4()}",
                        'name': f.name,
                        'source': 'local',
                        'mime_type': f.type,
                        'bytes': f.getvalue()
                    })
                if rejected:
                    st.error(f"❌ Arquivo(s) rejeitado(s) — tamanho máximo {MAX_FILE_SIZE_MB} MB: {', '.join(rejected)}")
                if new_files:
                    st.session_state.active_context_files = [
                        x for x in st.session_state.active_context_files if x['source'] != 'local'
                    ] + new_files
                    st.success(f"✓ {len(new_files)} arquivo(s) carregado(s)")
        
        with tab_drive:
            if st.session_state.get("user_drive_service"):
                try:
                    folders = [{'id': 'root', 'name': 'Meu Drive'}]
                    folders.extend(list_drive_folders(st.session_state["user_drive_service"]))
                    folder_map = {f['name']: f['id'] for f in folders}
                    selected = st.selectbox("Pasta:", list(folder_map.keys()), key="drive_folder")
                    
                    files = get_processable_drive_files_in_folder(st.session_state["user_drive_service"], folder_map[selected])
                    if files:
                        selected_files = st.multiselect(
                            "Arquivos:",
                            files,
                            format_func=lambda x: f"{x['name']} ({MIME_TYPE_DISPLAY.get(x['mimeType'], 'Arquivo')})",
                            key="drive_files"
                        )
                        if selected_files:
                            st.session_state.active_context_files = [
                                x for x in st.session_state.active_context_files if x['source'] != 'drive'
                            ] + [{'id': f['id'], 'name': f['name'], 'source': 'drive', 'mime_type': f['mimeType']} for f in selected_files]
                    else:
                        st.info("Nenhum arquivo encontrado nesta pasta.")
                except Exception as e:
                    st.error(f"Erro: {e}")
            else:
                st.warning("Google Drive não conectado.")
    
    if st.session_state.active_context_files:
        st.markdown('<div class="active-files">', unsafe_allow_html=True)
        cols = st.columns([0.9, 0.1])
        with cols[0]:
            files_html = "".join([
                f'<span class="file-chip">{_get_material_icon_html("description")} {f["name"][:20]}</span>'
                for f in st.session_state.active_context_files
            ])
            st.markdown(f'<div class="active-files">{files_html}</div>', unsafe_allow_html=True)
        with cols[1]:
            if st.button("🗑️", key="clear_files", help="Limpar arquivos"):
                st.session_state.active_context_files = []
                st.rerun()
    
    chat_container = st.container(height=450, border=False)

    with st.form(key='chat_form', clear_on_submit=True):
        cols = st.columns([0.88, 0.12])
        with cols[0]:
            user_input = st.text_input(
                "Pergunta",
                placeholder="Faça qualquer pergunta sobre SST...",
                key="chat_input",
                label_visibility="collapsed"
            )
        with cols[1]:
            submitted = st.form_submit_button("", icon=":material/send:", use_container_width=True)

    query_to_process = None
    if submitted and user_input.strip():
        query_to_process = sanitize_text_input(user_input, max_length=5000)
    elif st.session_state.pending_query:
        query_to_process = st.session_state.pending_query
        st.session_state.pending_query = None

    if query_to_process:
        if not check_chat_quota():
            render_chat_quota_exceeded()
            with chat_container:
                if not st.session_state.messages:
                    _render_welcome_screen()
                else:
                    for idx, msg in enumerate(st.session_state.messages):
                        _render_message(msg, idx, markdown_func)
            return

        try:
            check_rate_limit("chat_llm")
        except RateLimitExceeded as rle:
            log_security_event(
                SecurityEvent.RATE_LIMIT_EXCEEDED,
                user_email=st.session_state.get("user_email"),
                feature="chat_llm",
                detail=str(rle),
            )
            st.warning(f"⏳ Muitas perguntas em um curto período. Aguarde {rle.retry_after:.0f} segundos e tente novamente.")
            with chat_container:
                if not st.session_state.messages:
                    _render_welcome_screen()
                else:
                    for idx, msg in enumerate(st.session_state.messages):
                        _render_message(msg, idx, markdown_func)
            return

        st.session_state.messages.append({"role": "user", "content": query_to_process})
        st.session_state.last_follow_ups = []

        _mode_instructions = {
            "quick": (
                "MODO CONSULTA RÁPIDA: Seja objetivo e conciso. "
                "Limite sua resposta a no máximo 3 parágrafos ou 10 itens de lista. "
                "Vá direto ao ponto sem introduções longas."
            ),
            "deep": (
                "MODO ANÁLISE TÉCNICA: Forneça análise aprofundada com referências normativas "
                "completas, incluindo artigos, subitens e anexos específicos das NRs quando relevante. "
                "Estruture a resposta com seções claras."
            ),
        }
        context_texts = [_mode_instructions.get(st.session_state.chat_mode, "")]
        for ctx_file in st.session_state.active_context_files:
            try:
                if ctx_file['source'] == 'local' and 'bytes' in ctx_file:
                    with tempfile.NamedTemporaryFile(delete=False, suffix=os.path.splitext(ctx_file['name'])[1]) as tmp:
                        tmp.write(ctx_file['bytes'])
                        tmp_path = tmp.name
                    text = get_text_from_file_path(tmp_path, ctx_file['name'], ctx_file['mime_type'])
                    os.unlink(tmp_path)
                    if text:
                        context_texts.append(f"[{ctx_file['name']}]: {text[:5000]}")
                elif ctx_file['source'] == 'drive':
                    file_bytes = get_file_bytes_by_id(st.session_state["user_drive_service"], ctx_file['id'], ctx_file['mime_type'])
                    if file_bytes:
                        ext = '.docx' if 'document' in ctx_file['mime_type'] else '.pdf'
                        with tempfile.NamedTemporaryFile(delete=False, suffix=ext) as tmp:
                            tmp.write(file_bytes)
                            tmp_path = tmp.name
                        text = get_text_from_file_path(tmp_path, ctx_file['name'], ctx_file['mime_type'])
                        os.unlink(tmp_path)
                        if text:
                            context_texts.append(f"[{ctx_file['name']}]: {text[:5000]}")
            except Exception as e:
                logger.error(f"Erro contexto {ctx_file['name']}: {e}")

        full_text = ""
        downloads = []

        drive_search_kw = _extract_drive_search_keyword(query_to_process)
        drive_search_results: list[dict] = []
        if drive_search_kw:
            try:
                drive_svc = st.session_state.get("user_drive_service") or "app_service"
                raw_files = list_drive_files_by_keyword(drive_svc, drive_search_kw, max_results=6)
                for f in raw_files:
                    if f.get("webViewLink"):
                        drive_search_results.append({
                            "document_name": f.get("name", "Arquivo"),
                            "file_type": f.get("mimeType", "application/octet-stream"),
                            "drive_file_id": f.get("id", ""),
                            "url_viewer": f.get("webViewLink", ""),
                        })
            except Exception as _e:
                logger.debug(f"[chat] Drive keyword search failed: {_e}")

        with chat_container:
            for idx, msg in enumerate(st.session_state.messages):
                _render_message(msg, idx, markdown_func)

            ai_icon = _get_material_icon_html("smart_toy")
            stream_slot = st.empty()
            stream_slot.markdown(
                f'<div class="typing-indicator">'
                f'<div class="msg-ai-avatar">{ai_icon}</div>'
                f'<div class="typing-dots"><span></span><span></span><span></span></div>'
                f'</div>',
                unsafe_allow_html=True,
            )

            try:
                qa_instance = NRQuestionAnswering(on_status=make_streamlit_status_callback())
                chat_history = [{"role": m["role"], "content": m["content"]} for m in st.session_state.messages[:-1]]
                stream_gen = qa_instance.stream_answer_question(query_to_process, chat_history, context_texts)

                chunk_buffer = ""
                render_counter = 0

                for chunk in stream_gen:
                    full_text += chunk
                    chunk_buffer += chunk
                    render_counter += 1
                    if render_counter % 6 == 0 or '\n' in chunk_buffer:
                        content_html = markdown_func(full_text + " ▌")
                        stream_slot.markdown(
                            f'<div class="msg-ai"><div class="msg-ai-avatar">{ai_icon}</div>'
                            f'<div class="msg-ai-content">{content_html}</div></div>',
                            unsafe_allow_html=True,
                        )
                        chunk_buffer = ""

                if full_text:
                    content_html = markdown_func(full_text)
                    stream_slot.markdown(
                        f'<div class="msg-ai"><div class="msg-ai-avatar">{ai_icon}</div>'
                        f'<div class="msg-ai-content">{content_html}</div></div>',
                        unsafe_allow_html=True,
                    )

                downloads = qa_instance.get_last_suggested_downloads()
                seen_ids = {d.get("drive_file_id") for d in downloads}
                for dr in drive_search_results:
                    if dr["drive_file_id"] not in seen_ids:
                        downloads.append(dr)
                        seen_ids.add(dr["drive_file_id"])
                if full_text:
                    follow_ups = _generate_follow_ups(query_to_process, full_text)
                    st.session_state.last_follow_ups = follow_ups

            except Exception as e:
                logger.error(f"Erro ao processar: {e}", exc_info=True)
                full_text = "Ocorreu um erro ao processar sua pergunta. Por favor, tente novamente."
                stream_slot.markdown(
                    f'<div class="msg-ai"><div class="msg-ai-avatar">{ai_icon}</div>'
                    f'<div class="msg-ai-content">{full_text}</div></div>',
                    unsafe_allow_html=True,
                )

        st.session_state.messages.append({
            "role": "assistant",
            "content": full_text or "Desculpe, não consegui processar sua pergunta.",
            "suggested_downloads": downloads,
        })
        increment_chat_messages_today()
        st.rerun()

    else:
        with chat_container:
            if not st.session_state.messages:
                _render_welcome_screen()
            else:
                for idx, msg in enumerate(st.session_state.messages):
                    _render_message(msg, idx, markdown_func)
                if st.session_state.last_follow_ups:
                    _render_follow_ups(st.session_state.last_follow_ups)
