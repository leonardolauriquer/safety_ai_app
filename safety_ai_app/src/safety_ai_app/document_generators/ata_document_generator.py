# src/safety_ai_app/document_generators/ata_document_generator.py

import logging
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from io import BytesIO
from datetime import date
from typing import Dict, Any, List, Optional
import os
import base64
from docxtpl import DocxTemplate

logger = logging.getLogger(__name__)

# Cores da paleta (verde neon e tons de cinza/preto)
ACCENT_COLOR = RGBColor(0x27, 0xAE, 0x60) # Verde neon
TEXT_COLOR = RGBColor(0xFF, 0xFF, 0xFF) # Branco para texto em fundo escuro
MEDIUM_GRAY_TEXT = RGBColor(0x8B, 0x94, 0x9E) # Cinza médio para texto secundário
BLACK_TEXT = RGBColor(0x00, 0x00, 0x00) # Preto para texto de assinatura e conteúdo de tabela

def _add_page_number_and_company_info(footer_obj, company_name="Safety AI App"):
    """Adiciona numeração de página e informações da empresa ao rodapé do documento."""
    # Limpa o rodapé existente para evitar duplicação
    for paragraph in list(footer_obj.paragraphs):
        footer_obj._element.remove(paragraph._element)
    for table in list(footer_obj.tables): # Limpa tabelas também
        footer_obj._element.remove(table._element)

    paragraph = footer_obj.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    run_company = paragraph.add_run(f"{company_name} | ")
    run_company.font.size = Pt(8)
    run_company.font.color.rgb = MEDIUM_GRAY_TEXT

    run_page = paragraph.add_run()
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    run_page._r.append(fldChar1)

    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = 'PAGE'
    run_page._r.append(instrText)

    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'end')
    run_page._r.append(fldChar2)
    run_page.font.size = Pt(8)
    run_page.font.color.rgb = MEDIUM_GRAY_TEXT

    run_of = paragraph.add_run(' de ')
    run_of.font.size = Pt(8)
    run_of.font.color.rgb = MEDIUM_GRAY_TEXT

    run_numpages = paragraph.add_run()
    fldChar3 = OxmlElement('w:fldChar')
    fldChar3.set(qn('w:fldCharType'), 'begin')
    run_numpages._r.append(fldChar3)

    instrText2 = OxmlElement('w:instrText')
    instrText2.set(qn('xml:space'), 'preserve')
    instrText2.text = 'NUMPAGES'
    run_numpages._r.append(instrText2)

    fldChar4 = OxmlElement('w:fldChar')
    fldChar4.set(qn('w:fldCharType'), 'end')
    run_numpages._r.append(fldChar4)
    run_numpages.font.size = Pt(8)
    run_numpages.font.color.rgb = MEDIUM_GRAY_TEXT

def _set_cell_borders(cell):
    """Define bordas para uma célula da tabela com espessura de 1.5pt."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for border_name in ['top', 'left', 'bottom', 'right']:
        border_elm = OxmlElement(f'w:{border_name}')
        border_elm.set(qn('w:val'), 'single')
        border_elm.set(qn('w:sz'), '12') # 1.5 pt (12/8 pt) para maior visibilidade
        border_elm.set(qn('w:color'), 'auto') # Cor automática para bordas
        tcBorders.append(border_elm)
    tcPr.append(tcBorders)

def _create_header_table_content(header_obj, ata_data: Dict[str, Any], user_logo_base64: Optional[str]) -> None:
    """Cria o conteúdo da tabela do cabeçalho com logo, que se repetirá em todas as páginas."""
    app_logo_path = os.path.join(os.path.dirname(__file__), '..', '..', 'assets', 'app_logo.png')

    header_table_total_width = Cm(3) + Cm(10) + Cm(5)
    header_table = header_obj.add_table(rows=1, cols=3, width=header_table_total_width)
    header_table.autofit = False
    header_table.columns[0].width = Cm(3)
    header_table.columns[1].width = Cm(10)
    header_table.columns[2].width = Cm(5)
    header_table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Célula 1: Logo
    logo_cell = header_table.cell(0, 0)
    try:
        if user_logo_base64:
            logo_bytes = base64.b64decode(user_logo_base64.split(',', 1)[1])
            image_stream = BytesIO(logo_bytes)
            logo_cell.paragraphs[0].add_run().add_picture(image_stream, width=Cm(2.5))
            logger.info("Logo do usuário adicionado ao cabeçalho da Ata.")
        elif os.path.exists(app_logo_path):
            logo_cell.paragraphs[0].add_run().add_picture(app_logo_path, width=Cm(2.5))
            logger.info("Logo da aplicação adicionado ao cabeçalho da Ata.")
        else:
            logo_cell.text = "LOGO"
            logo_cell.paragraphs[0].runs[0].font.size = Pt(8)
            logo_cell.paragraphs[0].runs[0].font.color.rgb = MEDIUM_GRAY_TEXT
            logger.warning("Nenhum logo encontrado ou fornecido para o cabeçalho da Ata.")
        logo_cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    except Exception as e:
        logger.error(f"Erro ao adicionar logo ao cabeçalho do documento da Ata: {e}")
        logo_cell.text = "LOGO"
        logo_cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        logo_cell.paragraphs[0].runs[0].font.size = Pt(8)
        logo_cell.paragraphs[0].runs[0].font.color.rgb = MEDIUM_GRAY_TEXT

    # Célula 2: Título Central
    title_cell = header_table.cell(0, 1)
    title_paragraph = title_cell.paragraphs[0]
    title_run = title_paragraph.add_run("ATA DE DDS / TREINAMENTO / REUNIÃO")
    title_run.font.size = Pt(16)
    title_run.font.bold = True
    title_run.font.color.rgb = ACCENT_COLOR
    title_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Célula 3: Informações da Ata
    info_cell = header_table.cell(0, 2)
    info_paragraph = info_cell.paragraphs[0]
    info_paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    info_paragraph.add_run(f"Tipo: {ata_data.get('event_type', 'N/A')}\n").font.size = Pt(8)
    info_paragraph.add_run(f"Data: {ata_data.get('date', date.today()).strftime('%d/%m/%Y')}\n").font.size = Pt(8)
    info_paragraph.add_run(f"Hora: {ata_data.get('start_time', 'N/A')} - {ata_data.get('end_time', 'N/A')}").font.size = Pt(8)
    for run in info_paragraph.runs:
        run.font.color.rgb = MEDIUM_GRAY_TEXT


def _create_instructor_signature_content(doc_obj: Document, ata_data: Dict[str, Any]) -> List[Any]:
    """Cria o conteúdo do bloco de assinatura do instrutor e retorna como uma lista de elementos."""
    elements = []
    
    p = doc_obj.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    if ata_data.get('instructor_signature_image_base64'):
        try:
            signature_bytes = base64.b64decode(ata_data['instructor_signature_image_base64'].split(',', 1)[1])
            image_stream = BytesIO(signature_bytes)
            p.add_run().add_picture(image_stream, width=Cm(5))
            p.add_run("\nAssinatura Digital do Instrutor").font.size = Pt(10)
        except Exception as e:
            logger.error(f"Erro ao adicionar assinatura digital do instrutor: {e}")
            p.add_run("_________________________________________").font.size = Pt(10)
            p.add_run("\nAssinatura do Instrutor (Erro ao carregar digital)").font.size = Pt(10)
    else:
        p.add_run("_________________________________________").font.size = Pt(10)
        p.add_run("\nAssinatura do Instrutor").font.size = Pt(10)
    
    p.add_run(f"\n{ata_data.get('instructor_name', 'N/A')}").font.size = Pt(10)
    p.add_run("\nInstrutor / Responsável").font.size = Pt(10)

    for run in p.runs:
        run.font.color.rgb = BLACK_TEXT
    
    elements.append(p._element)
    return elements

def _create_participants_table_content(doc_obj: Document, participants_data: List[Dict[str, Any]]) -> List[Any]:
    """
    Cria a tabela de participantes com nome, CPF e assinatura.
    """
    elements = []

    if not participants_data:
        p = doc_obj.add_paragraph("Nenhum participante informado.", style='Normal')
        p.paragraph_format.space_after = Pt(10)
        elements.append(p._element)
        return elements

    table = doc_obj.add_table(rows=1, cols=3)
    table.autofit = False
    table.allow_autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    col_widths = [Cm(6), Cm(4), Cm(7)] # Nome, CPF, Assinatura
    for i, width in enumerate(col_widths):
        table.columns[i].width = width
    
    hdr_cells = table.rows[0].cells
    headers = ['Nome Completo', 'CPF', 'Assinatura']
    for i, header_text in enumerate(headers):
        cell = hdr_cells[i]
        cell.text = header_text
        cell.paragraphs[0].runs[0].font.bold = True
        cell.paragraphs[0].runs[0].font.color.rgb = BLACK_TEXT
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        _set_cell_borders(cell)

    for i, participant in enumerate(participants_data):
        row_cells = table.add_row().cells
        row_cells[0].text = participant.get('name', 'N/A')
        row_cells[1].text = participant.get('cpf', 'N/A')
        
        # Assinatura
        signature_cell = row_cells[2]
        signature_cell.vertical_alignment = WD_ALIGN_VERTICAL.BOTTOM
        p_sig = signature_cell.paragraphs[0]
        p_sig.alignment = WD_ALIGN_PARAGRAPH.CENTER

        if participant.get('signature_image_base64'):
            try:
                signature_bytes = base64.b64decode(participant['signature_image_base64'].split(',', 1)[1])
                image_stream = BytesIO(signature_bytes)
                p_sig.add_run().add_picture(image_stream, width=Cm(4)) # Assinatura menor para participantes
            except Exception as e:
                logger.error(f"Erro ao adicionar assinatura digital para {participant.get('name')}: {e}")
                p_sig.add_run("_________________________").font.size = Pt(9)
        else:
            p_sig.add_run("_________________________").font.size = Pt(9)
        
        for cell in row_cells:
            cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
            cell.paragraphs[0].runs[0].font.size = Pt(9)
            cell.paragraphs[0].runs[0].font.color.rgb = BLACK_TEXT
            _set_cell_borders(cell)

    elements.append(table._element)
    return elements

def _create_attachments_content(doc_obj: Document, attachments_data: List[Dict[str, Any]]) -> List[Any]:
    """
    Cria o conteúdo dos anexos, incorporando fotos e notas de texto, e listando documentos.
    """
    elements = []

    if not attachments_data:
        p = doc_obj.add_paragraph("Nenhum anexo informado.", style='Normal')
        p.paragraph_format.space_after = Pt(10)
        elements.append(p._element)
        return elements

    for i, attachment in enumerate(attachments_data):
        p_desc = doc_obj.add_paragraph(f"Anexo {i+1}: {attachment.get('description', 'N/A')}", style='Normal')
        p_desc.runs[0].font.bold = True
        p_desc.runs[0].font.color.rgb = BLACK_TEXT
        elements.append(p_desc._element)

        if attachment.get('type') == "Nota de Texto" and attachment.get('file_base64'):
            p_text = doc_obj.add_paragraph(attachment['file_base64'], style='Normal')
            p_text.paragraph_format.space_after = Pt(5)
            elements.append(p_text._element)
        elif attachment.get('type') == "Foto" and attachment.get('file_base64'):
            try:
                img_bytes = base64.b64decode(attachment['file_base64'].split(',', 1)[1])
                img_stream = BytesIO(img_bytes)
                p_img = doc_obj.add_paragraph()
                p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p_img.add_run().add_picture(img_stream, width=Cm(10)) # Tamanho padrão para fotos
                p_img.paragraph_format.space_after = Pt(5)
                elements.append(p_img._element)
            except Exception as e:
                logger.error(f"Erro ao adicionar foto de anexo '{attachment.get('file_name')}': {e}")
                p_error = doc_obj.add_paragraph(f"[Erro ao carregar foto: {attachment.get('file_name')}]")
                elements.append(p_error._element)
        elif attachment.get('type') == "Documento (PDF/DOCX/TXT)" and attachment.get('file_name'):
            p_doc = doc_obj.add_paragraph(f"Documento: {attachment['file_name']} (Tipo: {attachment.get('file_type', 'N/A')})", style='Normal')
            p_doc.paragraph_format.space_after = Pt(5)
            elements.append(p_doc._element)
        
        elements.append(doc_obj.add_paragraph("")._element) # Espaço entre anexos

    return elements


def _insert_content_at_placeholder(doc_obj: Document, placeholder_text: str, content_creation_func, *args, **kwargs) -> bool:
    """
    Encontra um parágrafo com o texto do placeholder e insere conteúdo gerado por content_creation_func
    no lugar do placeholder.
    """
    for p in doc_obj.paragraphs:
        if placeholder_text in p.text:
            parent_element = p._element.getparent()
            idx = parent_element.index(p._element)
            parent_element.remove(p._element)

            content_elements = content_creation_func(doc_obj, *args, **kwargs)
            
            for element in content_elements:
                parent_element.insert(idx, element)
                idx += 1
            return True
    return False

def create_ata_document(ata_data: Dict[str, Any], user_logo_base64: Optional[str] = None) -> BytesIO:
    """
    Cria um documento DOCX de Ata com base nos dados fornecidos, utilizando um template DOCX.
    """
    script_dir = os.path.dirname(__file__)
    template_path = os.path.join(script_dir, 'templates', 'ata_template.docx')

    if not os.path.exists(template_path):
        logger.error(f"Template DOCX não encontrado em: {template_path}. Usando fallback de geração sem template.")
        # Fallback simples, pode ser expandido se necessário
        doc = Document()
        doc.add_heading('Ata de DDS / Treinamento / Reunião', level=1)
        doc.add_paragraph(f"Tipo: {ata_data.get('event_type', 'N/A')}")
        doc.add_paragraph(f"Título: {ata_data.get('title', 'N/A')}")
        doc.add_paragraph(f"Data: {ata_data.get('date', date.today()).strftime('%d/%m/%Y')}")
        doc.add_paragraph(f"Local: {ata_data.get('location', 'N/A')}")
        doc.add_paragraph(f"Instrutor: {ata_data.get('instructor_name', 'N/A')}")
        doc.add_paragraph(f"Conteúdo: {ata_data.get('content', 'N/A')}")
        doc.add_heading('Participantes', level=2)
        for p in ata_data.get('participants', []):
            doc.add_paragraph(f"- {p.get('name', 'N/A')} (CPF: {p.get('cpf', 'N/A')})")
        doc.add_heading('Assinatura do Instrutor', level=2)
        doc.add_paragraph("_________________________")
        doc.add_paragraph(ata_data.get('instructor_name', 'N/A'))
        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer

    doc = DocxTemplate(template_path)

    context = {
        'event_type': ata_data.get('event_type', 'N/A'),
        'title': ata_data.get('title', 'N/A'),
        'date_formatted': ata_data.get('date', date.today()).strftime('%d/%m/%Y'),
        'start_time': ata_data.get('start_time', 'N/A'),
        'end_time': ata_data.get('end_time', 'N/A'),
        'location': ata_data.get('location', 'N/A'),
        'instructor_name': ata_data.get('instructor_name', 'N/A'),
        'content': ata_data.get('content', 'N/A'),
        'instructor_signature_placeholder': '[[INSTRUCTOR_SIGNATURE_BLOCK]]',
        'participants_table_placeholder': '[[PARTICIPANTS_TABLE]]',
        'attachments_block_placeholder': '[[ATTACHMENTS_BLOCK]]',
    }
    
    try:
        doc.render(context)
    except Exception as e:
        logger.error(f"Erro ao renderizar template DOCX da Ata com docxtpl: {e}", exc_info=True)
        raise

    rendered_buffer = BytesIO()
    doc.save(rendered_buffer)
    rendered_buffer.seek(0)
    final_document = Document(rendered_buffer)

    # --- Configurar cabeçalhos e rodapés ---
    section = final_document.sections[0]
    section.is_different_first_page = False # Logo em todas as páginas

    def _clear_header_footer(hf_obj):
        for p in list(hf_obj.paragraphs):
            hf_obj._element.remove(p._element)
        for tbl in list(hf_obj.tables):
            hf_obj._element.remove(tbl._element)

    if section.header:
        _clear_header_footer(section.header)
    if section.footer:
        _clear_header_footer(section.footer)

    _create_header_table_content(section.header, ata_data, user_logo_base64)
    _add_page_number_and_company_info(section.footer)

    # --- Inserir o bloco de assinatura do instrutor ---
    _insert_content_at_placeholder(
        final_document, 
        '[[INSTRUCTOR_SIGNATURE_BLOCK]]', 
        _create_instructor_signature_content, 
        ata_data
    )

    # --- Inserir a Tabela de Participantes ---
    _insert_content_at_placeholder(
        final_document,
        '[[PARTICIPANTS_TABLE]]',
        _create_participants_table_content,
        ata_data.get('participants', [])
    )

    # --- Inserir o Bloco de Anexos ---
    _insert_content_at_placeholder(
        final_document,
        '[[ATTACHMENTS_BLOCK]]',
        _create_attachments_content,
        ata_data.get('attachments', [])
    )

    buffer = BytesIO()
    final_document.save(buffer)
    buffer.seek(0)
    return buffer
