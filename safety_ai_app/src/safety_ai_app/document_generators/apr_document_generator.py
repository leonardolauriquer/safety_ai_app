# src/safety_ai_app/document_generators/apr_document_generator.py

import logging
from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm, Mm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from io import BytesIO
from datetime import date
from typing import Dict, Any, List, Optional
import os
import base64
from docxtpl import DocxTemplate

logger = logging.getLogger(__name__)

# Cores da paleta (verde neon e tons de cinza/preto)
ACCENT_COLOR = RGBColor(0x27, 0xAE, 0x60) # Verde neon
TEXT_COLOR = RGBColor(0xFF, 0xFF, 0xFF) # Branco para texto em fundo escuro (usado para cabeçalhos com fundo escuro)
MEDIUM_GRAY_TEXT = RGBColor(0x8B, 0x94, 0x9E) # Cinza médio para texto secundário
BLACK_TEXT = RGBColor(0x00, 0x00, 0x00) # Preto para texto de assinatura e conteúdo de tabela

def _add_page_number_and_company_info(footer_obj, company_name="Safety AI App"):
    """Adiciona numeração de página e informações da empresa ao rodapé do documento."""
    # Limpa o rodapé existente para evitar duplicação
    for paragraph in list(footer_obj.paragraphs):
        footer_obj._element.remove(paragraph._element)
    for table in list(footer_obj.tables): # Limpa tabelas também
        footer_obj._element.remove(table._element)

    # Adicionar um novo parágrafo ao rodapé após a limpeza
    paragraph = footer_obj.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Adiciona informações da empresa à esquerda
    run_company = paragraph.add_run(f"{company_name} | ")
    run_company.font.size = Pt(8)
    run_company.font.color.rgb = MEDIUM_GRAY_TEXT

    # Adiciona o campo de número de página
    run_page = paragraph.add_run()
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    run_page._r.append(fldChar1)

    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = 'PAGE'
    run_page._r.append(instrText)

    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'end')
    run_page._r.append(fldChar2)
    run_page.font.size = Pt(8)
    run_page.font.color.rgb = MEDIUM_GRAY_TEXT

    # Adiciona " de "
    run_of = paragraph.add_run(' de ')
    run_of.font.size = Pt(8)
    run_of.font.color.rgb = MEDIUM_GRAY_TEXT

    # Adiciona o campo de número total de páginas
    run_numpages = paragraph.add_run()
    fldChar3 = OxmlElement('w:fldChar')
    fldChar3.set(qn('w:fldCharType'), 'begin')
    run_numpages._r.append(fldChar3)

    instrText2 = OxmlElement('w:instrText')
    instrText2.set(qn('xml:space'), 'preserve')
    instrText2.text = 'NUMPAGES'
    run_numpages._r.append(instrText2)

    fldChar4 = OxmlElement('w:fldChar')
    fldChar4.set(qn('w:fldCharType'), 'end')
    run_numpages._r.append(fldChar4)
    run_numpages.font.size = Pt(8)
    run_numpages.font.color.rgb = MEDIUM_GRAY_TEXT

def _set_cell_borders(cell):
    """Define bordas para uma célula da tabela com espessura de 1.5pt."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for border_name in ['top', 'left', 'bottom', 'right']:
        border_elm = OxmlElement(f'w:{border_name}')
        border_elm.set(qn('w:val'), 'single')
        border_elm.set(qn('w:sz'), '12') # 1.5 pt (12/8 pt) para maior visibilidade
        border_elm.set(qn('w:color'), 'auto') # Cor automática para bordas
        tcBorders.append(border_elm)
    tcPr.append(tcBorders)

def _create_header_table_content(header_obj, apr_data: Dict[str, Any], user_logo_base64: Optional[str], include_logo: bool) -> None:
    """Cria o conteúdo da tabela do cabeçalho (com ou sem logo)."""
    app_logo_path = os.path.join(os.path.dirname(__file__), '..', '..', 'assets', 'app_logo.png')

    header_table_total_width = Cm(3) + Cm(10) + Cm(5) # Soma das larguras das colunas
    header_table = header_obj.add_table(rows=1, cols=3, width=header_table_total_width)
    header_table.autofit = False
    header_table.columns[0].width = Cm(3)
    header_table.columns[1].width = Cm(10)
    header_table.columns[2].width = Cm(5)
    header_table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Célula 1: Logo (apenas se include_logo for True)
    logo_cell = header_table.cell(0, 0)
    if include_logo: # This is the key check
        try:
            if user_logo_base64: # Prioriza o logo do usuário se fornecido
                logo_bytes = base64.b64decode(user_logo_base64.split(',', 1)[1]) # Pega apenas a parte base64
                image_stream = BytesIO(logo_bytes)
                logo_cell.paragraphs[0].add_run().add_picture(image_stream, width=Cm(2.5))
                logger.info("Logo do usuário adicionado ao cabeçalho.")
            elif os.path.exists(app_logo_path): # Fallback para o logo da aplicação
                logo_cell.paragraphs[0].add_run().add_picture(app_logo_path, width=Cm(2.5))
                logger.info("Logo da aplicação adicionado ao cabeçalho.")
            else:
                logo_cell.text = "LOGO"
                logo_cell.paragraphs[0].runs[0].font.size = Pt(8)
                logo_cell.paragraphs[0].runs[0].font.color.rgb = MEDIUM_GRAY_TEXT
                logger.warning("Nenhum logo encontrado ou fornecido para o cabeçalho.")
            logo_cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        except Exception as e:
            logger.error(f"Erro ao adicionar logo ao cabeçalho do documento: {e}")
            logo_cell.text = "LOGO"
            logo_cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            logo_cell.paragraphs[0].runs[0].font.size = Pt(8)
            logo_cell.paragraphs[0].runs[0].font.color.rgb = MEDIUM_GRAY_TEXT
    else:
        # Se não incluir logo, a célula pode ficar vazia ou com algum texto placeholder
        logo_cell.text = "" # Célula vazia para manter o layout

    # Célula 2: Título Central
    title_cell = header_table.cell(0, 1)
    title_paragraph = title_cell.paragraphs[0]
    title_run = title_paragraph.add_run("ANÁLISE PRELIMINAR DE RISCO (APR)")
    title_run.font.size = Pt(16)
    title_run.font.bold = True
    title_run.font.color.rgb = ACCENT_COLOR
    title_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Célula 3: Informações da APR
    info_cell = header_table.cell(0, 2)
    info_paragraph = info_cell.paragraphs[0]
    info_paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    info_paragraph.add_run(f"FRM - {apr_data.get('apr_number', 'XXXX')}\n").font.size = Pt(8)
    info_paragraph.add_run(f"REV.: {apr_data.get('revision_number', 'XX')}\n").font.size = Pt(8)
    info_paragraph.add_run(f"Data: {apr_data.get('start_date', date.today()).strftime('%d/%m/%Y')}").font.size = Pt(8)
    for run in info_paragraph.runs:
        run.font.color.rgb = MEDIUM_GRAY_TEXT


def _create_supervisor_signature_content(doc_obj: Document, apr_data: Dict[str, Any]) -> List[Any]:
    """Cria o conteúdo do bloco de assinatura do supervisor e retorna como uma lista de elementos."""
    elements = []
    
    p = doc_obj.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    if apr_data.get('supervisor_signature_image_base64'):
        try:
            signature_bytes = base64.b64decode(apr_data['supervisor_signature_image_base64'].split(',', 1)[1])
            image_stream = BytesIO(signature_bytes)
            p.add_run().add_picture(image_stream, width=Cm(5))
            p.add_run("\nAssinatura Digital do Supervisor").font.size = Pt(10)
        except Exception as e:
            logger.error(f"Erro ao adicionar assinatura digital do supervisor: {e}")
            p.add_run("_________________________________________").font.size = Pt(10)
            p.add_run("\nAssinatura do Supervisor (Erro ao carregar digital)").font.size = Pt(10)
    else:
        p.add_run("_________________________________________").font.size = Pt(10)
        p.add_run("\nAssinatura do Supervisor").font.size = Pt(10)
    
    p.add_run(f"\n{apr_data.get('supervisor', 'N/A')}").font.size = Pt(10)
    p.add_run("\nSupervisor / Encarregado Responsável").font.size = Pt(10)

    for run in p.runs:
        run.font.color.rgb = BLACK_TEXT # Cor preta para o texto da assinatura
    
    elements.append(p._element)
    return elements

def _create_approvers_signature_table_content(doc_obj: Document, approvers_data: List[Dict[str, Any]]) -> List[Any]:
    """Cria o conteúdo da tabela de assinaturas dos aprovadores e retorna como uma lista de elementos."""
    elements = []

    if not approvers_data:
        p = doc_obj.add_paragraph("Nenhum responsável pela aprovação informado.", style='Normal')
        elements.append(p._element)
        return elements

    table = doc_obj.add_table(rows=1, cols=len(approvers_data))
    table.autofit = False
    table.allow_autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    col_width = Cm(18 / len(approvers_data)) # Distribui a largura uniformemente
    for col in table.columns:
        col.width = col_width

    for i, approver in enumerate(approvers_data):
        cell = table.cell(0, i)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.BOTTOM
        
        p_cell = cell.paragraphs[0]
        p_cell.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        if approver.get('signature_image_base64'):
            try:
                signature_bytes = base64.b64decode(approver['signature_image_base64'].split(',', 1)[1])
                image_stream = BytesIO(signature_bytes)
                p_cell.add_run().add_picture(image_stream, width=Cm(5)) # Ajustado para Cm(5)
                p_cell.add_run("\n") # Adiciona quebra de linha após a imagem
            except Exception as e:
                logger.error(f"Erro ao adicionar assinatura digital para {approver.get('name')}: {e}")
                p_cell.add_run("_________________________").font.size = Pt(9)
                p_cell.add_run("\n")
        else:
            p_cell.add_run("_________________________").font.size = Pt(9)
            p_cell.add_run("\n")
        
        p_cell.add_run(f"{approver.get('name', 'N/A')}").font.size = Pt(9)
        p_cell.add_run(f"\n{approver.get('role', 'N/A')}").font.size = Pt(9)
        
        for run in p_cell.runs:
            run.font.color.rgb = BLACK_TEXT # Cor preta para o texto da assinatura
        # REMOVIDA A BORDA DA CÉLULA DE ASSINATURA
        # _set_cell_borders(cell) 
    
    elements.append(table._element)
    return elements

def _create_activity_steps_table_content(doc_obj: Document, activity_steps_data: List[Dict[str, Any]]) -> List[Any]:
    """
    Cria a tabela de detalhamento da atividade e análise de risco usando python-docx.
    Retorna como uma lista de elementos para inserção.
    """
    elements = []

    if not activity_steps_data: # Removido o check de 'Nenhuma etapa detalhada.' para não adicionar parágrafo extra
        p = doc_obj.add_paragraph("Nenhuma etapa de atividade detalhada.", style='Normal')
        p.paragraph_format.space_after = Pt(10) # Adiciona um pouco de espaço
        elements.append(p._element)
        return elements

    table = doc_obj.add_table(rows=1, cols=6)
    table.autofit = False
    table.allow_autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Ajuste das larguras das colunas para a tabela de atividades
    # Total da largura da tabela: ~16.5 cm (considerando margens A4)
    col_widths = [Cm(2.5), Cm(3.5), Cm(3.0), Cm(1.5), Cm(1.5), Cm(4.5)] # Total 16.5 Cm
    for i, width in enumerate(col_widths):
        table.columns[i].width = width
    
    # Cabeçalhos da tabela
    hdr_cells = table.rows[0].cells
    headers = ['Etapa da Atividade', 'Perigos e Riscos', 'Consequências', 'Probabilidade', 'Severidade', 'Medidas de Controle']
    for i, header_text in enumerate(headers):
        cell = hdr_cells[i]
        cell.text = header_text
        cell.paragraphs[0].runs[0].font.bold = True
        cell.paragraphs[0].runs[0].font.color.rgb = BLACK_TEXT # Alterado para PRETO
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        _set_cell_borders(cell)


    # Preencher as linhas da tabela com os dados
    for i, step in enumerate(activity_steps_data):
        row_cells = table.add_row().cells
        row_cells[0].text = f"Etapa {i+1}: {step.get('step_description', 'N/A')}"
        row_cells[1].text = step.get('hazards', 'N/A')
        row_cells[2].text = step.get('consequences', 'N/A')
        row_cells[3].text = step.get('probability', 'N/A')
        row_cells[4].text = step.get('severity', 'N/A')
        row_cells[5].text = step.get('controls', 'N/A')
        for cell in row_cells:
            cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT # Conteúdo alinhado à esquerda
            cell.paragraphs[0].runs[0].font.size = Pt(9)
            cell.paragraphs[0].runs[0].font.color.rgb = BLACK_TEXT # Alterado para PRETO
            _set_cell_borders(cell)

    elements.append(table._element)
    return elements

def _create_legend_table_content(doc_obj: Document) -> List[Any]:
    """
    Cria a tabela de Legenda usando python-docx.
    Retorna como uma lista de elementos para inserção.
    """
    elements = []

    legend_data = [
        ("APR", "Análise Preliminar de Riscos"),
        ("AVC", "Acidente Vascular Cerebral"),
        ("DDS", "Diálogo Diário de Segurança"),
        ("EPI", "Equipamento de Proteção Individual"),
        ("ETA", "Estação de Tratamento de Água"),
        ("ETE", "Estação de Tratamento de Efluentes"),
        ("FISPQ", "Ficha de Informação de Segurança do Produto Químico"),
        ("NR", "Norma Regulamentadora"),
        ("PFF", "Peça Facial Filtrante"),
        ("PQ", "Produto Químico"),
        ("PT", "Permissão de Trabalho"),
        ("PTA", "Plataforma de Trabalho Aéreo")
    ]

    table = doc_obj.add_table(rows=1, cols=2)
    table.autofit = False
    table.allow_autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT # Alinha a tabela à esquerda

    table.columns[0].width = Cm(3) # Largura para o Termo
    table.columns[1].width = Cm(14) # Largura para a Definição
    
    # Cabeçalhos da tabela
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Termo'
    hdr_cells[1].text = 'Definição'
    
    for cell in hdr_cells:
        cell.paragraphs[0].runs[0].font.bold = True
        cell.paragraphs[0].runs[0].font.color.rgb = BLACK_TEXT # Alterado para PRETO
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        _set_cell_borders(cell)

    # Preencher as linhas da tabela com os dados da legenda
    for termo, definicao in legend_data:
        row_cells = table.add_row().cells
        row_cells[0].text = termo
        row_cells[1].text = definicao
        for cell in row_cells:
            cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT # Conteúdo alinhado à esquerda
            cell.paragraphs[0].runs[0].font.size = Pt(9)
            cell.paragraphs[0].runs[0].font.color.rgb = BLACK_TEXT # Alterado para PRETO
            _set_cell_borders(cell)

    elements.append(table._element)
    return elements


def _insert_content_at_placeholder(doc_obj: Document, placeholder_text: str, content_creation_func, *args, **kwargs) -> bool:
    """
    Encontra um parágrafo com o texto do placeholder e insere conteúdo gerado por content_creation_func
    no lugar do placeholder.
    """
    for p in doc_obj.paragraphs:
        if placeholder_text in p.text:
            # Remove o parágrafo do placeholder
            parent_element = p._element.getparent()
            idx = parent_element.index(p._element)
            parent_element.remove(p._element)

            # Gera o conteúdo (lista de elementos XML)
            # A função content_creation_func deve receber doc_obj como primeiro argumento
            content_elements = content_creation_func(doc_obj, *args, **kwargs)
            
            # Insere cada elemento gerado no local do placeholder
            for element in content_elements:
                parent_element.insert(idx, element)
                idx += 1 # Incrementa o índice para inserir o próximo elemento após o anterior
            return True
    return False

def create_apr_document(apr_data: Dict[str, Any], user_logo_base64: Optional[str] = None) -> BytesIO:
    """
    Cria um documento DOCX de Análise Preliminar de Risco (APR) com base nos dados fornecidos,
    utilizando um template DOCX.
    """
    script_dir = os.path.dirname(__file__)
    template_path = os.path.join(script_dir, 'templates', 'apr_template.docx')
    app_logo_path = os.path.join(script_dir, '..', '..', 'assets', 'app_logo.png')

    if not os.path.exists(template_path):
        logger.error(f"Template DOCX não encontrado em: {template_path}. Usando fallback de geração sem template.")
        return _create_apr_document_without_template(apr_data, user_logo_base64)

    doc = DocxTemplate(template_path)

    # --- Preparar o contexto para o template ---
    context = {
        'apr_number': apr_data.get('apr_number', 'N/A'),
        'revision_number': apr_data.get('revision_number', 'N/A'),
        'start_date_formatted': apr_data.get('start_date', date.today()).strftime('%d/%m/%Y'),
        'end_date_formatted': apr_data['end_date'].strftime('%d/%m/%Y') if apr_data.get('end_date') else 'N/A',
        'work_schedule': apr_data.get('work_schedule', 'N/A'),
        'location': apr_data.get('location', 'N/A'),
        'company': apr_data.get('company', 'N/A'),
        'supervisor': apr_data.get('supervisor', 'N/A'),
        'task_name': apr_data.get('task_name', 'N/A'),
        'task_objective': apr_data.get('task_objective', 'N/A'),
        'other_epis': apr_data.get('other_epis', ''), # Alterado para string vazia para melhor condicional
        'other_equipments_tools': apr_data.get('other_equipments_tools', ''), # Alterado para string vazia
        'general_observations': apr_data.get('general_observations', 'N/A'),
        'emergency_contacts': apr_data.get('emergency_contacts', 'N/A'),
        'other_waste_disposal': apr_data.get('other_waste_disposal', ''), # Alterado para string vazia
        'other_additional_measures': apr_data.get('other_additional_measures', ''), # Alterado para string vazia
        'other_trainings': apr_data.get('other_trainings', ''), # Alterado para string vazia
        # Placeholders LITERAIS para blocos complexos que serão substituídos no pós-processamento
        'supervisor_signature_placeholder': '[[SUPERVISOR_SIGNATURE_BLOCK]]',
        'approvers_signature_placeholder': '[[APPROVERS_SIGNATURES_BLOCK]]',
        'activity_steps_placeholder': '[[ACTIVITY_STEPS_TABLE]]',
        'legend_table_placeholder': '[[LEGEND_TABLE]]', # Novo placeholder para a legenda
    }

    # --- Listas (usando bullet points) ---
    context['epis_list'] = [{'item': epi} for epi in apr_data.get('selected_epis', [])]
    if not context['epis_list']:
        context['epis_list'].append({'item': "Nenhum EPI selecionado/informado."})

    context['equipments_tools_list'] = [{'item': tool} for tool in apr_data.get('selected_equipments_tools', [])]
    if not context['equipments_tools_list']:
        context['equipments_tools_list'].append({'item': "Nenhum equipamento/ferramenta selecionado/informado."})

    context['waste_disposal_list'] = [{'item': waste} for waste in apr_data.get('waste_disposal', [])]
    if not context['waste_disposal_list']:
        context['waste_disposal_list'].append({'item': "Nenhuma informação sobre disposição de resíduos."})

    context['additional_measures_list'] = [{'item': measure} for measure in apr_data.get('additional_measures', [])]
    if not context['additional_measures_list']:
        context['additional_measures_list'].append({'item': "Nenhuma medida adicional informada."})

    context['trainings_list'] = [{'item': training} for training in apr_data.get('trainings', [])]
    if not context['trainings_list']:
        context['trainings_list'].append({'item': "Nenhum treinamento específico informado."})
    
    # --- Renderizar o documento com o contexto ---
    try:
        doc.render(context)
    except Exception as e:
        logger.error(f"Erro ao renderizar template DOCX com docxtpl: {e}", exc_info=True)
        raise

    # --- Pós-processamento para cabeçalho, rodapé e tabelas complexas ---
    rendered_buffer = BytesIO()
    doc.save(rendered_buffer)
    rendered_buffer.seek(0)
    final_document = Document(rendered_buffer)

    # --- Configurar cabeçalhos e rodapés ---
    section = final_document.sections[0]
    # Desativar a opção de primeira página diferente para que o cabeçalho padrão se aplique a todas
    section.is_different_first_page = False 

    # Limpa cabeçalhos e rodapés existentes para substituí-los de forma mais robusta
    def _clear_header_footer(hf_obj):
        for p in list(hf_obj.paragraphs):
            hf_obj._element.remove(p._element)
        for tbl in list(hf_obj.tables):
            hf_obj._element.remove(tbl._element)

    # Limpa o cabeçalho e rodapé padrão
    if section.header:
        _clear_header_footer(section.header)
    if section.footer:
        _clear_header_footer(section.footer)

    # Cria o cabeçalho padrão (com logo, que agora se aplica a todas as páginas)
    _create_header_table_content(section.header, apr_data, user_logo_base64, include_logo=True)
    
    # Adicionar rodapé padrão
    _add_page_number_and_company_info(section.footer)

    # --- Inserir o bloco de assinatura do supervisor ---
    _insert_content_at_placeholder(
        final_document, 
        '[[SUPERVISOR_SIGNATURE_BLOCK]]', 
        _create_supervisor_signature_content, 
        apr_data
    )

    # --- Inserir a Tabela de Detalhamento da Atividade e Análise de Risco ---
    _insert_content_at_placeholder(
        final_document,
        '[[ACTIVITY_STEPS_TABLE]]',
        _create_activity_steps_table_content,
        apr_data.get('activity_steps', [])
    )

    # --- Inserir a Tabela de Legenda ---
    _insert_content_at_placeholder(
        final_document,
        '[[LEGEND_TABLE]]',
        _create_legend_table_content
    )

    # --- Inserir a tabela de assinaturas dos aprovadores ---
    _insert_content_at_placeholder(
        final_document, 
        '[[APPROVERS_SIGNATURES_BLOCK]]', 
        _create_approvers_signature_table_content, 
        apr_data.get('approvers', [])
    )

    # Salvar o documento final em um buffer de bytes
    buffer = BytesIO()
    final_document.save(buffer)
    buffer.seek(0)
    return buffer

# Fallback function if template is not found
def _create_apr_document_without_template(apr_data: Dict[str, Any], user_logo_base64: Optional[str] = None) -> BytesIO:
    """
    Cria um documento DOCX de Análise Preliminar de Risco (APR) sem usar um template,
    como fallback caso o template não seja encontrado.
    (Este é o código anterior, adaptado para ser um fallback)
    """
    document = Document()

    # --- Configurações de Estilo Global ---
    style = document.styles['Normal']
    font = style.font
    font.name = 'Inter'
    font.size = Pt(11)

    # Adicionar estilos customizados
    h1_style = document.styles['Heading 1']
    h1_font = h1_style.font
    h1_font.name = 'Inter'
    h1_font.size = Pt(14)
    h1_font.bold = True
    h1_font.color.rgb = ACCENT_COLOR
    h1_style.paragraph_format.space_before = Pt(12)
    h1_style.paragraph_format.space_after = Pt(6)

    h2_style = document.styles['Heading 2']
    h2_font = h2_style.font
    h2_font.name = 'Inter'
    h2_font.size = Pt(12)
    h2_font.bold = True
    h2_font.color.rgb = MEDIUM_GRAY_TEXT
    h2_style.paragraph_format.space_before = Pt(10)
    h2_style.paragraph_format.space_after = Pt(5)

    normal_style = document.styles['Normal']
    normal_font = normal_style.font
    normal_font.name = 'Inter'
    normal_font.size = Pt(10)
    normal_font.color.rgb = MEDIUM_GRAY_TEXT
    normal_style.paragraph_format.space_before = Pt(0)
    normal_style.paragraph_format.space_after = Pt(3)
    normal_style.paragraph_format.line_spacing = 1.15

    # No fallback, o cabeçalho e rodapé são mais simples e se repetem
    # Criar o cabeçalho padrão (com logo, pois é um fallback mais simples)
    _create_header_table_content(document.sections[0].header, apr_data, user_logo_base64, include_logo=True)
    _add_page_number_and_company_info(document.sections[0].footer)

    document.add_paragraph("")

    document.add_heading('1. Identificação da APR', level=1)
    document.add_paragraph(f"**Número da APR:** {apr_data.get('apr_number', 'N/A')}")
    document.add_paragraph(f"**Revisão:** {apr_data.get('revision_number', 'N/A')}")
    document.add_paragraph(f"**Data de Início:** {apr_data.get('start_date', date.today()).strftime('%d/%m/%Y')}")
    if apr_data.get('end_date'):
        document.add_paragraph(f"**Data de Término:** {apr_data['end_date'].strftime('%d/%m/%Y')}")
    document.add_paragraph(f"**Horário de Trabalho:** {apr_data.get('work_schedule', 'N/A')}")
    document.add_paragraph(f"**Local / Área / Setor:** {apr_data.get('location', 'N/A')}")
    document.add_paragraph(f"**Empresa / Contratada:** {apr_data.get('company', 'N/A')}")
    
    document.add_paragraph(f"**Supervisor / Encarregado Responsável:** {apr_data.get('supervisor', 'N/A')}")
    if apr_data.get('supervisor_signature_image_base64'):
        try:
            signature_bytes = base64.b64decode(apr_data['supervisor_signature_image_base64'].split(',', 1)[1])
            image_stream = BytesIO(signature_bytes)
            document.add_picture(image_stream, width=Cm(5))
            document.add_paragraph("Assinatura Digital do Supervisor")
        except Exception as e:
            logger.error(f"Erro ao adicionar assinatura digital do supervisor: {e}")
            document.add_paragraph("_________________________________________")
            document.add_paragraph("Assinatura do Supervisor (Erro ao carregar digital)")
    else:
        document.add_paragraph("_________________________________________")
        document.add_paragraph("Assinatura do Supervisor")

    document.add_paragraph("")

    document.add_heading('2. Detalhes da Atividade/Tarefa', level=1)
    document.add_paragraph(f"**Nome da Atividade / Tarefa Específica:** {apr_data.get('task_name', 'N/A')}")
    document.add_paragraph(f"**Objetivo da Atividade:** {apr_data.get('task_objective', 'N/A')}")

    document.add_paragraph("")

    document.add_heading('3. EPIs e Equipamentos/Ferramentas a Serem Usados', level=1)
    
    document.add_paragraph("**EPIs Necessários:**")
    if apr_data.get('selected_epis'):
        for epi in apr_data['selected_epis']:
            document.add_paragraph(f"- {epi}", style='List Bullet')
    if apr_data.get('other_epis'):
        for epi in apr_data['other_epis'].split('\n'):
            if epi.strip():
                document.add_paragraph(f"- {epi.strip()}", style='List Bullet')
    if not apr_data.get('selected_epis') and not apr_data.get('other_epis'):
        document.add_paragraph("- Nenhum EPI selecionado/informado.")

    document.add_paragraph("")

    document.add_paragraph("**Equipamentos e Ferramentas Necessários:**")
    if apr_data.get('selected_equipments_tools'):
        for tool in apr_data['selected_equipments_tools']:
            document.add_paragraph(f"- {tool}", style='List Bullet')
    if apr_data.get('other_equipments_tools'):
        for tool in apr_data['other_equipments_tools'].split('\n'):
            if tool.strip():
                document.add_paragraph(f"- {tool.strip()}", style='List Bullet')
    if not apr_data.get('selected_equipments_tools') and not apr_data.get('other_equipments_tools'):
        document.add_paragraph("- Nenhum equipamento/ferramenta selecionado/informado.")

    document.add_paragraph("")

    document.add_heading('4. Detalhamento da Atividade e Análise de Risco', level=1)
    # Inserir a tabela de etapas da atividade usando a nova função helper
    activity_steps_elements = _create_activity_steps_table_content(document, apr_data.get('activity_steps', []))
    for element in activity_steps_elements:
        document._body._element.append(element)
        
    document.add_paragraph("")

    document.add_paragraph("")

    document.add_heading('5. Observações Gerais', level=1)
    document.add_paragraph(apr_data.get('general_observations', 'N/A'))

    document.add_paragraph("")

    document.add_heading('6. Contatos de Emergência', level=1)
    document.add_paragraph(apr_data.get('emergency_contacts', 'N/A'))

    document.add_paragraph("")

    document.add_heading('7. Disposição de Resíduos', level=1)
    if apr_data.get('waste_disposal'):
        for waste in apr_data['waste_disposal']:
            document.add_paragraph(f"- {waste}", style='List Bullet')
    if apr_data.get('other_waste_disposal'):
        for waste in apr_data['other_waste_disposal'].split('\n'):
            if waste.strip():
                document.add_paragraph(f"- {waste.strip()}", style='List Bullet')
    if not apr_data.get('waste_disposal') and not apr_data.get('other_waste_disposal'):
        document.add_paragraph("- Nenhum informação sobre disposição de resíduos.")

    document.add_paragraph("")

    document.add_heading('8. Medidas a Serem Adotadas pelos Envolvidos na Execução da Tarefa', level=1)
    if apr_data.get('additional_measures'):
        for measure in apr_data['additional_measures']:
            document.add_paragraph(f"- {measure}", style='List Bullet')
    else:
        document.add_paragraph("- Nenhuma medida adicional informada.")
    document.add_paragraph("")

    document.add_heading('9. Legenda', level=1) # Adicionado de volta
    # No fallback, inserimos o conteúdo da legenda diretamente.
    legend_elements = _create_legend_table_content(document)
    for element in legend_elements:
        document._body._element.append(element)
    document.add_paragraph("") # Adiciona um parágrafo vazio após a tabela da legenda

    document.add_heading('10. Treinamentos', level=1)
    if apr_data.get('trainings'):
        for training in apr_data['trainings']:
            document.add_paragraph(f"- {training}", style='List Bullet')
    else:
        document.add_paragraph("- Nenhum treinamento específico informado.")
    document.add_paragraph("")

    document.add_heading('11. Responsáveis pela Aprovação', level=1)
    # No fallback, inserimos o conteúdo da tabela de aprovadores diretamente.
    approvers_elements = _create_approvers_signature_table_content(document, apr_data.get('approvers', []))
    for element in approvers_elements:
        document._body._element.append(element)

    buffer = BytesIO()
    document.save(buffer)
    buffer.seek(0)
    return buffer